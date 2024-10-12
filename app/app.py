from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
import os
import streamlit as st

# Função para formatar títulos e texto em ABNT
def formatar_abnt(titulo, texto):
    doc = Document()

    # Configura título (centralizado e em negrito)
    par_titulo = doc.add_paragraph()
    titulo_run = par_titulo.add_run(titulo)
    titulo_run.font.size = Pt(14)
    titulo_run.bold = True
    titulo_run.font.name = 'Arial'
    par_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Configura o texto com espaçamento 1.5
    paragrafo = doc.add_paragraph(texto)
    run = paragrafo.runs[0]
    run.font.size = Pt(12)
    run.font.name = 'Arial'
    paragrafo.paragraph_format.line_spacing = 1.5

    # Alinhamento justificado
    paragrafo.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # Salva o documento formatado
    doc.save('texto_abnt.docx')
    print("Documento salvo como texto_abnt.docx")

st.title("Formatador de Textos ABNT")
st.markdown("Formate o seu texto no Padrão ABNT")
st.markdown(" Importante apertar 'Enter' nos campos")

# Exemplo de uso
titulo = st.text_input("Digite o seu título")
titulo_formatado = titulo.upper()
texto = st.text_area(label="Digite o seu texto")

if titulo and texto:
    formatar_abnt(titulo_formatado, texto)

    with open("texto_abnt.docx", "rb") as file:
        doc_bytes = file.read()

    # Botão de download
    send = st.download_button(
        label="Baixar Documento ABNT",
        data=doc_bytes,  # Passa o conteúdo do arquivo como bytes
        file_name="texto_abnt.docx",  # Nome do arquivo
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    # Apaga o arquivo após o download
    if send:
        os.remove("texto_abnt.docx")
        st.success("O arquivo foi baixado e resetado com sucesso.")
